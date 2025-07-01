import streamlit as st
import pandas as pd
import numpy as np
import datetime
import requests
import os
from datetime import datetime, time
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
import plotly.colors as pc
import plotly.io as pio
import io
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# DeepSeek API configuration
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-8fa69e0815904d0daf8831374c9999b8")
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"
MODEL = "deepseek-chat"

# Sidebar settings
st.sidebar.title("Settings")
min_speed = st.sidebar.number_input("Minimum Speed (RPM)", value=10)
idle_time = st.sidebar.number_input("Idle Time Threshold (min)", value=10)
user_machine_map = st.sidebar.text_area(
    "User-Machine Mapping (user:machine1,machine2)",
    value="user1:Extruder-80A\nuser2:Extruder-90B"
)

# Multiple file upload and selection
st.title("IoT Machine Analytics Dashboard (Multi-file, DeepSeek AI)")
uploaded_files = st.file_uploader("Upload machine CSV files", type="csv", accept_multiple_files=True)

if not uploaded_files:
    st.info("Please upload one or more CSV files to continue.")
    st.stop()

file_names = [file.name for file in uploaded_files]
selected_file = st.sidebar.selectbox("Select a file for analysis", file_names)

# Load the selected file
for file in uploaded_files:
    if file.name == selected_file:
        df = pd.read_csv(file)
        df.columns = [c.strip().lower() for c in df.columns]
        df['timestamp'] = pd.to_datetime(df['timestamp'], errors='coerce')
        df = df.dropna(subset=['timestamp']).sort_values('timestamp')
        break

# Derive running/idle status
df['is_running'] = df['speed'] >= min_speed
df['idle'] = ~df['is_running']

# Add time-based columns for analysis
df['date'] = df['timestamp'].dt.date
df['time'] = df['timestamp'].dt.time
df['hour'] = df['timestamp'].dt.hour

# INITIALIZE VARIABLES EARLY - BEFORE THEY ARE USED
min_date = df['timestamp'].dt.date.min()
max_date = df['timestamp'].dt.date.max()

# Set default values for date_range and time_range
date_range = (min_date, max_date)
time_range = (0, 23)

# KPI calculations (ensure these are calculated on the full df initially for overall summary)
total_minutes = (df['timestamp'].iloc[-1] - df['timestamp'].iloc[0]).total_seconds() / 60 if len(df) > 1 else 0
running_minutes = df['is_running'].sum()
uptime_percent = 100 * running_minutes / total_minutes if total_minutes else 0
avg_speed = df.loc[df['is_running'], 'speed'].mean() if running_minutes else 0
total_production = df['quantity'].iloc[-1] - df['quantity'].iloc[0] if len(df) > 1 else 0

df['idle_group'] = (df['idle'] != df['idle'].shift()).cumsum()
idle_events = df[df['idle']].groupby('idle_group').agg(
    start=('timestamp', 'first'),
    end=('timestamp', 'last'),
    duration=('timestamp', lambda x: (x.iloc[-1] - x.iloc[0]).total_seconds() / 60)
)

long_idles = idle_events[idle_events['duration'] >= idle_time]

# Table of low speed instances
df['low_speed'] = df['speed'] < min_speed
df['low_speed_group'] = (df['low_speed'] != df['low_speed'].shift()).cumsum()
low_speed_events = df[df['low_speed']].groupby('low_speed_group').agg(
    start=('timestamp', 'first'),
    end=('timestamp', 'last'),
    duration=('timestamp', lambda x: round((x.iloc[-1] - x.iloc[0]).total_seconds() / 60, 4)),
    min_speed=('speed', 'min'),
    avg_speed=('speed', 'mean')
).reset_index(drop=True)

st.subheader("Low Speed Events (Speed < Minimum Threshold)")
st.dataframe(low_speed_events)

# Dashboard display
st.subheader("Key Performance Indicators")
kpi1, kpi2, kpi3, kpi4 = st.columns(4)
kpi1.metric("Uptime (%)", f"{uptime_percent:.1f}")
kpi2.metric("Avg Speed (RPM)", f"{avg_speed:.1f}")
kpi3.metric("Production", f"{total_production:.1f}")
kpi4.metric(f"Idle Events (≥{idle_time} min)", f"{len(long_idles)}")

# Enhanced Graph Section with Time Filtering and View Options
st.subheader("Machine Performance Analysis")

# Create columns for controls
col1, col2, col3, col4 = st.columns(4)

with col1:
    graph_view = st.selectbox(
        "Select Graph View",
        ["Total Graph", "Running Only", "Idle Only", "Low Speed Events", "Production vs Speed"]
    )

with col2:
    chart_type = st.selectbox(
        "Select Chart Type",
        ["Bar Chart", "Line Chart", "Area Chart"]
    )

with col3:
    date_range = st.date_input(
        "Select Date Range",
        value=(min_date, max_date),
        min_value=min_date,
        max_value=max_date,
        format="YYYY/MM/DD"
    )

with col4:
    time_range = st.slider(
        "Select Time Range (Hours)",
        min_value=0,
        max_value=23,
        value=(0, 23),
        step=1,
        format="%d:00"
    )

# Filter data based on selections
if len(date_range) == 2:
    start_date, end_date = date_range
    filtered_df = df[
        (df['timestamp'].dt.date >= start_date) & 
        (df['timestamp'].dt.date <= end_date) &
        (df['hour'] >= time_range[0]) &
        (df['hour'] <= time_range[1])
    ].copy()
else:
    filtered_df = df[
        (df['timestamp'].dt.date == date_range[0]) &
        (df['hour'] >= time_range[0]) &
        (df['hour'] <= time_range[1])
    ].copy()

# Display filtered data info
st.info(f"Showing data from {filtered_df['timestamp'].min()} to {filtered_df['timestamp'].max()} ({len(filtered_df)} records)")

# Function to create bar charts using Plotly
def create_bar_chart(data, title, color='blue'):
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=data.index,
        y=data['speed'],
        name='Speed (RPM)',
        marker_color=color,
        text=data['speed'].round(1),
        textposition='outside'
    ))
    
    fig.update_layout(
        title=title,
        xaxis_title="Time",
        yaxis_title="Speed (RPM)",
        showlegend=True,
        height=500,
        xaxis=dict(tickangle=45)
    )
    
    return fig

# Function to create multi-series bar charts
def create_multi_bar_chart(data, columns, title):
    fig = go.Figure()
    
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
    
    for i, col in enumerate(columns):
        fig.add_trace(go.Bar(
            x=data.index,
            y=data[col],
            name=col,
            marker_color=colors[i % len(colors)]
        ))
    
    fig.update_layout(
        title=title,
        xaxis_title="Time",
        yaxis_title="Value",
        barmode='group',
        showlegend=True,
        height=500,
        xaxis=dict(tickangle=45)
    )
    
    return fig

# Create different graph views based on selection
if graph_view == "Total Graph":
    st.subheader("Speed Over Time - Complete View")
    
    if chart_type == "Bar Chart":
        chart_data = filtered_df.set_index('timestamp')[['speed']]
        fig = create_bar_chart(chart_data, "Machine Speed Over Time", '#1f77b4')
        st.plotly_chart(fig, use_container_width=True)
    elif chart_type == "Area Chart":
        chart_data = filtered_df.set_index('timestamp')[['speed']]
        st.area_chart(chart_data, use_container_width=True)
    else:
        chart_data = filtered_df.set_index('timestamp')[['speed']]
        st.line_chart(chart_data, use_container_width=True)
    
    filtered_uptime = (filtered_df['is_running'].sum() / len(filtered_df)) * 100 if len(filtered_df) > 0 else 0
    filtered_avg_speed = filtered_df['speed'].mean() if len(filtered_df) > 0 else 0
    st.write(f"**Filtered Period Stats:** Uptime: {filtered_uptime:.1f}%, Avg Speed: {filtered_avg_speed:.1f} RPM")

elif graph_view == "Running Only":
    st.subheader("Speed Over Time - Running Periods Only")
    running_data = filtered_df[filtered_df['is_running']].copy()
    if len(running_data) > 0:
        if chart_type == "Bar Chart":
            chart_data = running_data.set_index('timestamp')[['speed']]
            fig = create_bar_chart(chart_data, "Machine Speed - Running Periods Only", '#00ff00')
            st.plotly_chart(fig, use_container_width=True)
        elif chart_type == "Area Chart":
            chart_data = running_data.set_index('timestamp')[['speed']]
            st.area_chart(chart_data, use_container_width=True, color='#00ff00')
        else:
            chart_data = running_data.set_index('timestamp')[['speed']]
            st.line_chart(chart_data, use_container_width=True, color='#00ff00')
        
        st.write(f"**Running Periods:** {len(running_data)} records, Avg Speed: {running_data['speed'].mean():.1f} RPM")
    else:
        st.warning("No running periods found in the selected time range.")

elif graph_view == "Idle Only":
    st.subheader("Speed Over Time - Idle Periods Only")
    idle_data = filtered_df[filtered_df['idle']].copy()
    if len(idle_data) > 0:
        if chart_type == "Bar Chart":
            chart_data = idle_data.set_index('timestamp')[['speed']]
            fig = create_bar_chart(chart_data, "Machine Speed - Idle Periods Only", '#ff0000')
            st.plotly_chart(fig, use_container_width=True)
        elif chart_type == "Area Chart":
            chart_data = idle_data.set_index('timestamp')[['speed']]
            st.area_chart(chart_data, use_container_width=True, color='#ff0000')
        else:
            chart_data = idle_data.set_index('timestamp')[['speed']]
            st.line_chart(chart_data, use_container_width=True, color='#ff0000')
        
        st.write(f"**Idle Periods:** {len(idle_data)} records, Avg Speed: {idle_data['speed'].mean():.1f} RPM")
    else:
        st.warning("No idle periods found in the selected time range.")

elif graph_view == "Low Speed Events":
    st.subheader("Speed Over Time - Low Speed Events Highlighted")
    
    if chart_type == "Bar Chart":
        fig = go.Figure()
        
        normal_data = filtered_df[~filtered_df['low_speed']]
        if len(normal_data) > 0:
            fig.add_trace(go.Bar(
                x=normal_data['timestamp'],
                y=normal_data['speed'],
                name='Normal Speed',
                marker_color='#1f77b4'
            ))
        
        low_speed_data = filtered_df[filtered_df['low_speed']]
        if len(low_speed_data) > 0:
            fig.add_trace(go.Bar(
                x=low_speed_data['timestamp'],
                y=low_speed_data['speed'],
                name='Low Speed',
                marker_color='#ff0000'
            ))
        
        fig.update_layout(
            title="Machine Speed - Low Speed Events Highlighted",
            xaxis_title="Time",
            yaxis_title="Speed (RPM)",
            barmode='overlay',
            height=500,
            xaxis=dict(tickangle=45)
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
    else:
        chart_data = filtered_df.set_index('timestamp')[['speed']].copy()
        low_speed_data = filtered_df[filtered_df['low_speed']].set_index('timestamp')[['speed']].copy()
        
        if len(low_speed_data) > 0:
            chart_data.columns = ['Normal Speed']
            low_speed_data.columns = ['Low Speed']
            combined_data = pd.concat([chart_data, low_speed_data], axis=1)
            
            if chart_type == "Area Chart":
                st.area_chart(combined_data, use_container_width=True)
            else:
                st.line_chart(combined_data, use_container_width=True)
        else:
            if chart_type == "Area Chart":
                st.area_chart(chart_data, use_container_width=True)
            else:
                st.line_chart(chart_data, use_container_width=True)
    
    low_speed_count = len(filtered_df[filtered_df['low_speed']])
    st.write(f"**Low Speed Events:** {low_speed_count} records in selected period")

elif graph_view == "Production vs Speed":
    st.subheader("Production and Speed Correlation")
    filtered_df['production_rate'] = filtered_df['quantity'].diff().fillna(0)
    
    if chart_type == "Bar Chart":
        chart_data = filtered_df.set_index('timestamp')[['speed', 'production_rate']].copy()
        chart_data.columns = ['Speed (RPM)', 'Production Rate']
        fig = create_multi_bar_chart(chart_data, ['Speed (RPM)', 'Production Rate'], 
                                     "Production Rate vs Speed Comparison")
        st.plotly_chart(fig, use_container_width=True)
    else:
        chart_data = filtered_df.set_index('timestamp')[['speed', 'production_rate']].copy()
        chart_data.columns = ['Speed (RPM)', 'Production Rate']
        
        if chart_type == "Area Chart":
            st.area_chart(chart_data, use_container_width=True)
        else:
            st.line_chart(chart_data, use_container_width=True)
    
    if len(filtered_df) > 1:
        correlation = filtered_df['speed'].corr(filtered_df['production_rate'])
        st.write(f"**Speed-Production Correlation:** {correlation:.3f}")

# Production Over Time with Time Filtering
st.subheader("Production Over Time")

if chart_type == "Bar Chart":
    production_data = filtered_df.set_index('timestamp')[['quantity']]
    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=production_data.index,
        y=production_data['quantity'],
        name='Production Quantity',
        marker_color='#2ca02c'
    ))
    
    fig.update_layout(
        title="Production Quantity Over Time",
        xaxis_title="Time",
        yaxis_title="Quantity",
        height=400,
        xaxis=dict(tickangle=45)
    )
    
    st.plotly_chart(fig, use_container_width=True)
else:
    production_chart_data = filtered_df.set_index('timestamp')[['quantity']]
    if chart_type == "Area Chart":
        st.area_chart(production_chart_data, use_container_width=True)
    else:
        st.line_chart(production_chart_data, use_container_width=True)

# Enhanced Idle Events Table with Time Information and Graph
st.subheader("Idle Events Analysis")

# Filter idle events based on time selection
if len(date_range) == 2:
    start_datetime = datetime.combine(date_range[0], time(time_range[0], 0))
    end_datetime = datetime.combine(date_range[1], time(time_range[1], 59, 59)) # Fixed end time to include the whole hour
else:
    start_datetime = datetime.combine(date_range[0], time(time_range[0], 0))
    end_datetime = datetime.combine(date_range[0], time(time_range[1], 59, 59)) # Fixed end time to include the whole hour

# Ensure idle_events is not empty before filtering
if not idle_events.empty:
    filtered_idle_events = idle_events[
        (idle_events['start'] >= start_datetime) & 
        (idle_events['end'] <= end_datetime)
    ].copy()
else:
    filtered_idle_events = pd.DataFrame() # Initialize as empty if no idle events exist

if len(filtered_idle_events) > 0:
    filtered_idle_events['start_time'] = filtered_idle_events['start'].dt.strftime('%H:%M:%S')
    filtered_idle_events['end_time'] = filtered_idle_events['end'].dt.strftime('%H:%M:%S')
    filtered_idle_events['start_date'] = filtered_idle_events['start'].dt.date
    filtered_idle_events['idle_group'] = filtered_idle_events.index
    
    display_columns = ['idle_group', 'start_date', 'start_time', 'end_time', 'duration']
    st.write("**Idle Events Table:**")
    st.dataframe(filtered_idle_events[display_columns])
    
    st.write("**Idle Events Timeline Visualization:**")
    
    gantt_data = []
    for idx, row in filtered_idle_events.iterrows():
        gantt_data.append(dict(
            Task=f"Idle Group {idx}",
            Start=row['start'],
            Finish=row['end'],
            Duration=f"{row['duration']:.1f} min"
        ))
    
    if gantt_data:
        num_tasks = len(filtered_idle_events)
        base_colors = pc.qualitative.Plotly
        colors = (base_colors * ((num_tasks // len(base_colors)) + 1))[:num_tasks]
        
        try:
            fig = ff.create_gantt(
                gantt_data,
                colors=colors,
                index_col='Duration',
                title='Idle Events Timeline',
                show_colorbar=True,
                bar_width=0.3,
                showgrid_x=True,
                showgrid_y=True
            )
            
            fig.update_layout(
                height=400,
                xaxis_title="Time",
                yaxis_title="Idle Events",
                font=dict(size=10)
            )
            
            st.plotly_chart(fig, use_container_width=True)
            
        except Exception as e:
            st.warning(f"Could not create Gantt chart: {str(e)}")
            st.info("Showing alternative timeline visualization...")
            
            fig_alt = go.Figure()
            
            for idx, row in filtered_idle_events.iterrows():
                fig_alt.add_trace(go.Scatter(
                    x=[row['start'], row['end']],
                    y=[f"Group {idx}", f"Group {idx}"],
                    mode='lines+markers',
                    name=f"Idle Group {idx}",
                    line=dict(width=8),
                    marker=dict(size=10),
                    hovertemplate=f'<b>Group {idx}</b><br>' +
                                  f'Duration: %{row["duration"]:.1f} minutes<br>' +
                                  f'Start: %{x}<extra></extra>'
                ))
            
            fig_alt.update_layout(
                title="Idle Events Timeline (Alternative View)",
                xaxis_title="Time",
                yaxis_title="Idle Groups",
                height=400,
                showlegend=False
            )
            
            st.plotly_chart(fig_alt, use_container_width=True)
    
    st.write("**Idle Event Durations:**")
    
    fig_duration = go.Figure()
    fig_duration.add_trace(go.Bar(
        x=[f"Group {idx}" for idx in filtered_idle_events.index],
        y=filtered_idle_events['duration'],
        name='Duration (minutes)',
        marker_color='#ff6b6b',
        text=filtered_idle_events['duration'].round(1),
        textposition='outside',
        hovertemplate='<b>%{x}</b><br>' +
                      'Duration: %{y:.1f} minutes<br>' +
                      'Start: %{customdata[0]}<br>' +
                      'End: %{customdata[1]}<extra></extra>',
        customdata=list(zip(
            filtered_idle_events['start'].dt.strftime('%Y-%m-%d %H:%M:%S'),
            filtered_idle_events['end'].dt.strftime('%Y-%m-%d %H:%M:%S')
        ))
    ))
    
    fig_duration.update_layout(
        title="Idle Event Durations by Group",
        xaxis_title="Idle Group",
        yaxis_title="Duration (minutes)",
        height=400,
        showlegend=False
    )
    
    st.plotly_chart(fig_duration, use_container_width=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Idle Events", len(filtered_idle_events))
    with col2:
        st.metric("Total Idle Time", f"{filtered_idle_events['duration'].sum():.1f} min")
    with col3:
        st.metric("Average Idle Duration", f"{filtered_idle_events['duration'].mean():.1f} min")
    with col4:
        st.metric("Longest Idle Event", f"{filtered_idle_events['duration'].max():.1f} min")

else:
    st.info("No idle events found in the selected time range.")

# Hourly Performance Analysis
st.subheader("Hourly Performance Pattern")
hourly_stats = filtered_df.groupby('hour').agg({
    'speed': ['mean', 'max', 'min'],
    'is_running': 'mean',
    'quantity': lambda x: x.iloc[-1] - x.iloc[0] if len(x) > 1 else 0
}).round(2)

hourly_stats.columns = ['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Ratio', 'Production']
hourly_stats['Uptime_Percent'] = (hourly_stats['Uptime_Ratio'] * 100).round(1)

if chart_type == "Bar Chart":
    hourly_chart_data = hourly_stats[['Avg_Speed', 'Uptime_Percent']].copy()
    fig = create_multi_bar_chart(hourly_chart_data, ['Avg_Speed', 'Uptime_Percent'], 
                                 "Hourly Performance - Average Speed and Uptime")
    st.plotly_chart(fig, use_container_width=True)
else:
    hourly_chart_data = hourly_stats[['Avg_Speed', 'Uptime_Percent']].copy()
    if chart_type == "Area Chart":
        st.area_chart(hourly_chart_data, use_container_width=True)
    else:
        st.line_chart(hourly_chart_data, use_container_width=True)

st.write("**Hourly Performance Statistics:**")
st.dataframe(hourly_stats[['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Percent', 'Production']])

# DeepSeek API call function with error handling
def ask_deepseek(prompt, temperature=0.1):  # Lower temperature for more focused responses
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    # ENHANCED SYSTEM PROMPT FOR ACCURACY
    enhanced_system_prompt = """You are a senior IoT data analyst and procurement expert with 25+ years experience.

CRITICAL INSTRUCTIONS:
1. ALWAYS analyze the EXACT data provided in the string representations of the dataframes and summary statistics. Do NOT make assumptions.
2. Reference SPECIFIC values, durations, and patterns directly from the provided data summaries and tables.
3. Provide QUANTIFIED insights with precise numbers.
4. Focus on ACTIONABLE procurement and maintenance recommendations.
5. Include COST implications and ROI estimates where possible.
6. Verify your answer against the provided data before responding. If a specific detail is not present in the provided summarized data, state that clearly and do not hallucinate.

RESPONSE FORMAT:
- Start with key findings using actual data points from the provided data.
- Provide specific time periods or event IDs for patterns.
- Give concrete recommendations with part numbers/specifications (if applicable).
- Include maintenance scheduling with exact timeframes (if applicable).

EXPERTISE AREAS:
- Industrial machinery diagnostics
- Predictive maintenance strategies   
- Procurement cost optimization
- Equipment failure pattern analysis
- OEE (Overall Equipment Effectiveness) improvement
"""
    
    data = {
        "model": MODEL,
        "messages": [
            {"role": "system", "content": enhanced_system_prompt},
            {"role": "user", "content": prompt}
        ],
        "temperature": temperature,
        "max_tokens": 2000,  # Ensure detailed responses
        "top_p": 0.9  # Focus on most relevant responses
    }

    
    try:
        response = requests.post(DEEPSEEK_URL, headers=headers, json=data, timeout=300)
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            return f"DeepSeek API Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Error connecting to DeepSeek API: {str(e)}"

# Fixed: Function to create comprehensive data summary with explicit dataframe strings
def create_comprehensive_data_summary(dataframe, min_speed_threshold, idle_time_threshold):
    """
    Create a comprehensive summary of the dataset for DeepSeek analysis.
    This function processes the dataframe and returns structured summaries and
    string representations of relevant derived dataframes.
    """
    
    # Basic statistics
    total_duration_seconds = (dataframe['timestamp'].max() - dataframe['timestamp'].min()).total_seconds() if len(dataframe) > 1 else 0
    duration_hours = total_duration_seconds / 3600
    
    total_production_overall = dataframe['quantity'].iloc[-1] - dataframe['quantity'].iloc[0] if len(dataframe) > 1 else 0
    
    summary_stats = {
        'total_records': len(dataframe),
        'time_range_start': str(dataframe['timestamp'].min()),
        'time_range_end': str(dataframe['timestamp'].max()),
        'duration_hours': round(duration_hours, 2),
        'speed_stats': {
            'min': dataframe['speed'].min() if not dataframe['speed'].empty else 0,
            'max': dataframe['speed'].max() if not dataframe['speed'].empty else 0,
            'mean': dataframe['speed'].mean() if not dataframe['speed'].empty else 0
        },
        'quantity_stats': {
            'min': dataframe['quantity'].min() if not dataframe['quantity'].empty else 0,
            'max': dataframe['quantity'].max() if not dataframe['quantity'].empty else 0,
            'mean': dataframe['quantity'].mean() if not dataframe['quantity'].empty else 0,
            'total_production': total_production_overall
        }
    }
    
    # Operational Summary
    running_records = dataframe['is_running'].sum()
    idle_records = dataframe['idle'].sum()
    total_records_for_ops = len(dataframe)
    
    running_percentage = (running_records / total_records_for_ops * 100) if total_records_for_ops else 0
    idle_percentage = (idle_records / total_records_for_ops * 100) if total_records_for_ops else 0

    # Calculate long idle events from the passed dataframe context
    df_idle_groups = (dataframe['idle'] != dataframe['idle'].shift()).cumsum()
    df_idle_events = dataframe[dataframe['idle']].groupby(df_idle_groups).agg(
        start=('timestamp', 'first'),
        end=('timestamp', 'last'),
        duration=('timestamp', lambda x: (x.iloc[-1] - x.iloc[0]).total_seconds() / 60)
    )
    df_long_idles = df_idle_events[df_idle_events['duration'] >= idle_time_threshold]
    
    num_long_idle_events = len(df_long_idles)
    total_long_idle_duration = df_long_idles['duration'].sum() if not df_long_idles.empty else 0

    operational_summary = {
        'running_percentage': round(running_percentage, 1),
        'idle_percentage': round(idle_percentage, 1),
        'avg_running_speed': dataframe[dataframe['is_running']]['speed'].mean() if running_records else 0,
        'total_long_idle_events': num_long_idle_events,
        'total_long_idle_duration_minutes': round(total_long_idle_duration, 1)
    }

    # Hourly Performance (using the provided dataframe, which could be full or filtered)
    hourly_performance_df = dataframe.groupby('hour').agg({
        'speed': ['mean', 'max', 'min'],
        'is_running': 'mean',
        'quantity': lambda x: x.iloc[-1] - x.iloc[0] if len(x) > 1 else 0
    }).round(2)
    hourly_performance_df.columns = ['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Ratio', 'Production']
    hourly_performance_df['Uptime_Percent'] = (hourly_performance_df['Uptime_Ratio'] * 100).round(1)

    # Low Speed Events (using the provided dataframe, which could be full or filtered)
    df_low_speed = dataframe['speed'] < min_speed_threshold
    df_low_speed_groups = (df_low_speed != df_low_speed.shift()).cumsum()
    df_low_speed_events = dataframe[df_low_speed].groupby(df_low_speed_groups).agg(
        start=('timestamp', 'first'),
        end=('timestamp', 'last'),
        duration=('timestamp', lambda x: round((x.iloc[-1] - x.iloc[0]).total_seconds() / 60, 4)),
        min_speed=('speed', 'min'),
        avg_speed=('speed', 'mean')
    ).reset_index(drop=True)

    # Convert dataframes to string representations for the prompt
    hourly_performance_str = hourly_performance_df.to_string()
    low_speed_events_str = df_low_speed_events.to_string() if not df_low_speed_events.empty else "No low speed events found."
    long_idle_events_str = df_long_idles[['start', 'end', 'duration']].to_string() if not df_long_idles.empty else "No long idle events found (duration >= {} min).".format(idle_time_threshold)

    return summary_stats, operational_summary, hourly_performance_str, low_speed_events_str, long_idle_events_str

# AI-powered pattern analysis with token limits
st.subheader("AI-Powered Pattern Analysis (DeepSeek) - Dataframe Driven")
analysis_option = st.radio(
    "Select analysis scope:",
    ["Full Dataset Analysis", "Filtered Data Analysis", "Time Pattern Analysis"]
)

if st.button("🤖 Analyze Data with DeepSeek AI", type="primary"):
    
    # Determine which DataFrame to pass based on analysis_option
    data_for_analysis = df if analysis_option == "Full Dataset Analysis" else filtered_df

    if data_for_analysis.empty:
        st.warning("The selected dataset for analysis is empty. Please check your file or filters.")
        st.stop()

    summary_stats, operational_summary, hourly_performance_str, low_speed_events_str, long_idle_events_str = \
        create_comprehensive_data_summary(data_for_analysis, min_speed, idle_time)
    
    prompt_base = (
        f"**IoT MACHINE DATA ANALYSIS REPORT**\n\n"
        f"You are a senior IoT data analyst and procurement expert. Your task is to analyze the provided machine performance data and generate insights based *ONLY* on the data given in the summary statistics and detailed tables below. Do not make assumptions or infer data not explicitly present.\n\n"
        f"**1. Dataset Overview:**\n"
        f"- Total Records: {summary_stats['total_records']:,}\n"
        f"- Data Time Range: {summary_stats['time_range_start']} to {summary_stats['time_range_end']} ({summary_stats['duration_hours']} hours)\n"
        f"- Speed Range: {summary_stats['speed_stats']['min']:.1f}-{summary_stats['speed_stats']['max']:.1f} RPM (Average: {summary_stats['speed_stats']['mean']:.1f} RPM)\n"
        f"- Total Production: {summary_stats['quantity_stats']['total_production']:.1f} units\n\n"
        f"**2. Operational Summary:**\n"
        f"- Overall Uptime: {operational_summary['running_percentage']:.1f}%\n"
        f"- Overall Idle Time: {operational_summary['idle_percentage']:.1f}%\n"
        f"- Average Running Speed: {operational_summary['avg_running_speed']:.1f} RPM\n"
        f"- Total Critical Idle Events (duration >= {idle_time} min): {operational_summary['total_long_idle_events']} events, total duration: {operational_summary['total_long_idle_duration_minutes']:.1f} minutes\n\n"
        f"**3. Hourly Performance Statistics (Avg Speed, Max Speed, Min Speed, Uptime Percent, Production):**\n"
        f"```\n{hourly_performance_str}\n```\n\n"
        f"**4. Low Speed Events (Speed < {min_speed} RPM) Table:**\n"
        f"```\n{low_speed_events_str}\n```\n\n"
        f"**5. Long Idle Events (Duration >= {idle_time} min) Table:**\n"
        f"```\n{long_idle_events_str}\n```\n\n"
    )

    if analysis_option == "Full Dataset Analysis":
        prompt = prompt_base + (
            f"**Analysis Request for Full Dataset:**\n"
            f"Based on the provided full dataset summary and tables, provide:\n"
            f"1. Key performance patterns and anomalies observed over the entire period.\n"
            f"2. General maintenance recommendations based on overall trends and recurring issues.\n"
            f"3. Operational optimization suggestions for the machine's overall performance.\n"
            f"4. Broad procurement priorities derived from the data.\n"
            f"Ensure all insights are quantified with specific numbers and refer to the data provided above. If a particular insight is not directly derivable from the given data, state that limitation."
        )
        
    elif analysis_option == "Filtered Data Analysis":
        prompt = prompt_base + (
            f"**Analysis Request for Filtered Data:**\n"
            f"The above data represents a filtered view of the machine's performance. Based on this filtered dataset summary and tables, provide:\n"
            f"1. Key performance patterns and anomalies specific to this filtered time range.\n"
            f"2. Detailed maintenance recommendations relevant to the observed behavior in this period.\n"
            f"3. Operational optimization suggestions tailored to the conditions within this filtered period.\n"
            f"4. Procurement insights specific to the patterns identified in this filtered data.\n"
            f"Ensure all insights are highly specific, quantified with numbers, and directly supported by the provided filtered data. If a particular insight is not directly derivable from the given data, state that limitation."
        )
        
    else:  # Time Pattern Analysis
        # For Time Pattern Analysis, we want to focus on the hourly data specifically
        # and not necessarily the full tables of low speed/idle events, unless asked specifically.
        # We can reconstruct the prompt here to be more focused.
        hourly_stats_full_df = df.groupby('hour').agg({
            'speed': ['mean', 'std'],
            'is_running': 'mean',
            'quantity': lambda x: x.iloc[-1] - x.iloc[0] if len(x) > 1 else 0
        }).round(2)
        hourly_stats_full_df.columns = ['Avg_Speed', 'Speed_Std_Dev', 'Uptime_Ratio', 'Production']
        hourly_stats_full_df['Uptime_Percent'] = (hourly_stats_full_df['Uptime_Ratio'] * 100).round(1)
        hourly_performance_full_str = hourly_stats_full_df.to_string()

        prompt = (
            f"**TIME PATTERN ANALYSIS (Full Dataset Hourly Averages)**\n\n"
            f"You are a senior IoT data analyst and procurement expert. Your task is to analyze the provided hourly performance data and identify time-based patterns. Base your analysis *ONLY* on the data given below.\n\n"
            f"**Hourly Performance Statistics (Overall Dataset):**\n"
            f"```\n{hourly_performance_full_str}\n```\n\n"
            f"**Analysis Request:**\n"
            f"Analyze the hourly performance statistics to identify and quantify:\n"
            f"1. Peak performance hours and their characteristics (e.g., highest average speed, uptime, production).\n"
            f"2. Common low-performance or downtime periods (e.g., lowest uptime, average speed).\n"
            f"3. Specific maintenance scheduling recommendations based on optimal downtime windows or recurring low performance hours.\n"
            f"4. Any anomalies or significant deviations in hourly performance.\n"
            f"Ensure all findings are quantified with specific numbers and refer directly to the 'Hourly Performance Statistics' table. If a particular insight is not directly derivable from the given data, state that limitation."
        )
        
    with st.spinner("DeepSeek AI is analyzing your data..."):
        analysis = ask_deepseek(prompt)
        st.write("### 🤖 DeepSeek AI Analysis Results")
        st.write(analysis)

# FIXED: Enhanced User Query System, now also dataframe-driven
st.subheader("Ask DeepSeek AI About Your Data")
user_query = st.text_input("💬 Ask a question about your machine data")

if user_query:
    # Determine which DataFrame to pass for the user query context
    df_for_query_context = filtered_df if not filtered_df.empty else df
    
    if df_for_query_context.empty:
        st.warning("No data available to answer the query. Please upload a file or adjust filters.")
        st.stop()

    summary_stats_query, operational_summary_query, hourly_performance_query_str, \
    low_speed_events_query_str, long_idle_events_query_str = \
        create_comprehensive_data_summary(df_for_query_context, min_speed, idle_time)
    
    prompt = (
        f"**IoT MACHINE DATA QUERY AND ANALYSIS**\n\n"
        f"You are a senior IoT data analyst and procurement expert. Please answer the following question based *ONLY* on the provided summarized data. If the information is not directly available, state that you cannot answer from the provided context. Do not make assumptions.\n\n"
        f"**Context Data:**\n"
        f"**1. Dataset Overview:**\n"
        f"- Total Records: {summary_stats_query['total_records']:,}\n"
        f"- Data Time Range: {summary_stats_query['time_range_start']} to {summary_stats_query['time_range_end']} ({summary_stats_query['duration_hours']} hours)\n"
        f"- Average Speed: {summary_stats_query['speed_stats']['mean']:.1f} RPM\n"
        f"- Overall Uptime: {operational_summary_query['running_percentage']:.1f}%\n"
        f"- Total Production: {summary_stats_query['quantity_stats']['total_production']:.1f} units\n"
        f"- Long Idle Events (duration >= {idle_time} min): {operational_summary_query['total_long_idle_events']} events, total duration: {operational_summary_query['total_long_idle_duration_minutes']:.1f} minutes\n\n"
        f"**2. Hourly Performance Statistics:**\n"
        f"```\n{hourly_performance_query_str}\n```\n\n"
        f"**3. Low Speed Events Table:**\n"
        f"```\n{low_speed_events_query_str}\n```\n\n"
        f"**4. Long Idle Events Table:**\n"
        f"```\n{long_idle_events_query_str}\n```\n\n"
        f"**Question:** {user_query}\n\n"
        f"Please provide a specific and quantified answer, referencing the data points from the 'Context Data' above."
    )
    
    with st.spinner("🤖 DeepSeek AI is analyzing your question..."):
        answer = ask_deepseek(prompt)
        st.write("### 💡 DeepSeek AI Answer")
        st.write(answer)

# Advanced Analysis Options
st.subheader("Advanced AI Analysis Options")

col1, col2 = st.columns(2)

with col1:
    if st.button("🔧 Maintenance Schedule Recommendations"):
        # Always use full df for overall maintenance strategy
        summary_stats_full, operational_summary_full, hourly_performance_full_str, \
        low_speed_events_full_str, long_idle_events_full_str = \
            create_comprehensive_data_summary(df, min_speed, idle_time)
        
        maintenance_prompt = (
            f"**MAINTENANCE SCHEDULE ANALYSIS**\n\n"
            f"You are a senior IoT data analyst and procurement expert. Provide detailed maintenance recommendations and a schedule based *ONLY* on the provided data summaries and tables. Quantify everything and avoid external assumptions.\n\n"
            f"**Overall Dataset Context:**\n"
            f"- Total Records: {summary_stats_full['total_records']:,}, Data Period: {summary_stats_full['time_range_start']} to {summary_stats_full['time_range_end']}\n"
            f"- Overall Uptime: {operational_summary_full['running_percentage']:.1f}%, Total Long Idle Events: {operational_summary_full['total_long_idle_events']}\n\n"
            f"**Hourly Performance Statistics:**\n"
            f"```\n{hourly_performance_full_str}\n```\n\n"
            f"**Low Speed Events Table:**\n"
            f"```\n{low_speed_events_full_str}\n```\n\n"
            f"**Long Idle Events Table:**\n"
            f"```\n{long_idle_events_full_str}\n```\n\n"
            f"**Analysis Request:**\n"
            f"Analyze the above data to provide:\n"
            f"1. A predictive maintenance schedule based on identified patterns in low speed events, idle times, and hourly performance. Specify suggested timeframes (e.g., 'after X hours of operation' or 'during Y-Z AM').\n"
            f"2. Critical spare parts recommendations based on common failure indicators (e.g., prolonged low speed, frequent idle events) with justification. (If no specific part numbers are in the data, suggest types of parts).\n"
            f"3. Optimal maintenance windows (e.g., during consistent low production or idle periods).\n"
            f"4. Early warning indicators from the data that suggest impending issues or required maintenance.\n"
            f"5. A brief cost-benefit analysis for implementing these recommendations (e.g., 'reducing idle time by X% could save Y units of production').\n"
            f"All recommendations must be quantifiable and directly derivable from the provided data. State if data is insufficient for a specific point."
        )
        
        with st.spinner("🔧 Generating maintenance recommendations..."):
            maintenance_analysis = ask_deepseek(maintenance_prompt)
            st.write("### 🔧 Maintenance Recommendations")
            st.write(maintenance_analysis)

with col2:
    if st.button("📈 Production Optimization Analysis"):
        # Always use full df for overall optimization strategy
        summary_stats_full, operational_summary_full, hourly_performance_full_str, \
        low_speed_events_full_str, long_idle_events_full_str = \
            create_comprehensive_data_summary(df, min_speed, idle_time)
        
        optimization_prompt = (
            f"**PRODUCTION OPTIMIZATION ANALYSIS**\n\n"
            f"You are a senior IoT data analyst and procurement expert. Analyze the provided data summaries and tables to identify production optimization opportunities. Your analysis must be based *ONLY* on the provided data. Quantify all findings and recommendations.\n\n"
            f"**Overall Dataset Context:**\n"
            f"- Total Records: {summary_stats_full['total_records']:,}, Data Period: {summary_stats_full['time_range_start']} to {summary_stats_full['time_range_end']}\n"
            f"- Overall Uptime: {operational_summary_full['running_percentage']:.1f}%, Total Production: {summary_stats_full['quantity_stats']['total_production']:.1f} units\n\n"
            f"**Hourly Performance Statistics:**\n"
            f"```\n{hourly_performance_full_str}\n```\n\n"
            f"**Low Speed Events Table:**\n"
            f"```\n{low_speed_events_full_str}\n```\n\n"
            f"**Long Idle Events Table:**\n"
            f"```\n{long_idle_events_full_str}\n```\n\n"
            f"**Analysis Request:**\n"
            f"Analyze the above data for optimization opportunities:\n"
            f"1. Identify peak performance periods and the factors contributing to them (e.g., specific hours with high average speed and uptime).\n"
            f"2. Determine optimal operating parameters (e.g., ideal speed range) based on production output and uptime.\n"
            f"3. Suggest potential equipment upgrade considerations based on observed performance bottlenecks (e.g., recurring low speed, significant idle times that impact overall production). Quantify potential gains if possible.\n"
            f"4. Provide ROI calculations for suggested improvements (e.g., 'reducing idle events by X could increase production by Y units, generating Z revenue').\n"
            f"5. Outline procurement strategies to support optimized production (e.g., sourcing more reliable components, just-in-time spare part delivery for identified critical parts).\n"
            f"All suggestions must be data-backed and quantified. State explicitly if data is insufficient for a particular recommendation."
        )
        
        with st.spinner("📈 Analyzing production optimization opportunities..."):
            optimization_analysis = ask_deepseek(optimization_prompt)
            st.write("### 📈 Production Optimization")
            st.write(optimization_analysis)

# Enhanced function to create comprehensive Word report with graphs
def create_comprehensive_report_with_graphs(df, filtered_df, current_min_speed, current_idle_time, kpi_data, selected_file):
    doc = Document()
    
    title = doc.add_heading('IoT Machine Analytics Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('Executive Summary', level=1)
    
    summary_text = f"""
    This report provides a comprehensive analysis of machine performance data from {selected_file}.
    The analysis covers operational efficiency, idle events, and performance patterns to support
    procurement and maintenance decision-making.
    
    Key Findings:
    • Machine Uptime: {kpi_data['uptime']:.1f}%
    • Average Operating Speed: {kpi_data['avg_speed']:.1f} RPM
    • Total Production: {kpi_data['total_production']:.1f} units
    • Critical Idle Events (duration >= {current_idle_time} min): {kpi_data['idle_events']} events
    """
    
    doc.add_paragraph(summary_text)
    
    doc.add_heading('Dataset Overview', level=1)
    
    overview_table = doc.add_table(rows=1, cols=2)
    overview_table.style = 'Table Grid'
    
    hdr_cells = overview_table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'
    
    overview_data = [
        ('File Name', selected_file),
        ('Total Records', f"{len(df):,}"),
        ('Analysis Period', f"{df['timestamp'].min()} to {df['timestamp'].max()}"),
        ('Duration (Hours)', f"{((df['timestamp'].max() - df['timestamp'].min()).total_seconds() / 3600):.1f}"),
        ('Speed Range (RPM)', f"{df['speed'].min():.1f} - {df['speed'].max():.1f}"),
        ('Production Range', f"{df['quantity'].min():.1f} - {df['quantity'].max():.1f}")
    ]
    
    for param, value in overview_data:
        row_cells = overview_table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = str(value)
    
    doc.add_heading('Key Performance Indicators', level=1)
    
    kpi_table = doc.add_table(rows=1, cols=2)
    kpi_table.style = 'Table Grid'
    
    hdr_cells = kpi_table.rows[0].cells
    hdr_cells[0].text = 'KPI'
    hdr_cells[1].text = 'Value'
    
    kpi_rows = [
        ('Machine Uptime (%)', f"{kpi_data['uptime']:.1f}%"),
        ('Average Speed (RPM)', f"{kpi_data['avg_speed']:.1f}"),
        ('Total Production', f"{kpi_data['total_production']:.1f}"),
        ('Long Idle Events (duration >= {} min)'.format(current_idle_time), f"{kpi_data['idle_events']}")
    ]
    
    for kpi, value in kpi_rows:
        row_cells = kpi_table.add_row().cells
        row_cells[0].text = kpi
        row_cells[1].text = str(value)

    # --- Start of remaining original code from the user, ensuring the report generation is complete ---
    doc.add_heading('Machine Performance Analysis', level=1)

    # Add charts to the report - this requires saving them as images first
    # Example for "Total Graph - Line Chart"
    doc.add_heading('Speed Over Time - Complete View', level=2)
    chart_data = filtered_df.set_index('timestamp')[['speed']]
    fig_total = go.Figure()
    fig_total.add_trace(go.Scatter(x=chart_data.index, y=chart_data['speed'], mode='lines', name='Speed (RPM)'))
    fig_total.update_layout(title="Machine Speed Over Time", xaxis_title="Time", yaxis_title="Speed (RPM)")
    
    try:
        img_bytes = pio.to_image(fig_total, format="png")
    except Exception as e:
        doc.add_paragraph("Unable to add chart: " + str(e))
        return doc
    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
    doc.add_paragraph()

    # Add Production Over Time chart
    doc.add_heading('Production Over Time', level=2)
    production_chart_data = filtered_df.set_index('timestamp')[['quantity']]
    fig_prod = go.Figure()
    fig_prod.add_trace(go.Scatter(x=production_chart_data.index, y=production_chart_data['quantity'], mode='lines', name='Production Quantity'))
    fig_prod.update_layout(title="Production Quantity Over Time", xaxis_title="Time", yaxis_title="Quantity")
    
    img_bytes_prod = pio.to_image(fig_prod, format="png")
    doc.add_picture(io.BytesIO(img_bytes_prod), width=Inches(6))
    doc.add_paragraph()

    # Add Idle Events Analysis table and graph (if filtered_idle_events is not empty)
    doc.add_heading('Idle Events Analysis', level=1)
    
    # Re-calculate filtered_idle_events for the report to ensure it matches the current filtered_df context
    report_idle_groups = (filtered_df['idle'] != filtered_df['idle'].shift()).cumsum()
    report_idle_events = filtered_df[filtered_df['idle']].groupby(report_idle_groups).agg(
        start=('timestamp', 'first'),
        end=('timestamp', 'last'),
        duration=('timestamp', lambda x: (x.iloc[-1] - x.iloc[0]).total_seconds() / 60)
    )
    report_filtered_long_idles = report_idle_events[report_idle_events['duration'] >= current_idle_time]


    if not report_filtered_long_idles.empty:
        doc.add_paragraph("Idle Events Table (filtered by current selections):")
        # Convert DataFrame to a string table format for the report
        doc.add_paragraph(report_filtered_long_idles[['start', 'end', 'duration']].to_string())

        # Add Idle Events Timeline Visualization (if possible to generate Gantt chart)
        doc.add_paragraph("Idle Events Timeline Visualization:")
        gantt_data = []
        for idx, row in report_filtered_long_idles.iterrows():
            gantt_data.append(dict(
                Task=f"Idle Group {idx}",
                Start=row['start'],
                Finish=row['end'],
                Duration=f"{row['duration']:.1f} min"
            ))
        
        if gantt_data:
            num_tasks = len(report_filtered_long_idles)
            base_colors = pc.qualitative.Plotly
            colors = (base_colors * ((num_tasks // len(base_colors)) + 1))[:num_tasks]
            try:
                fig_gantt = ff.create_gantt(
                    gantt_data,
                    colors=colors,
                    index_col='Duration',
                    title='Idle Events Timeline',
                    show_colorbar=True,
                    bar_width=0.3,
                    showgrid_x=True,
                    showgrid_y=True
                )
                img_bytes_gantt = pio.to_image(fig_gantt, format="png")
                doc.add_picture(io.BytesIO(img_bytes_gantt), width=Inches(6))
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"Could not generate Gantt chart for the report: {e}")
                # Fallback for Gantt chart
                fig_alt = go.Figure()
                for idx, row in report_filtered_long_idles.iterrows():
                    fig_alt.add_trace(go.Scatter(
                        x=[row['start'], row['end']],
                        y=[f"Group {idx}", f"Group {idx}"],
                        mode='lines+markers',
                        name=f"Idle Group {idx}",
                        line=dict(width=8),
                        marker=dict(size=10)
                    ))
                fig_alt.update_layout(title="Idle Events Timeline (Alternative View)", xaxis_title="Time", yaxis_title="Idle Groups")
                img_bytes_alt = pio.to_image(fig_alt, format="png")
                doc.add_picture(io.BytesIO(img_bytes_alt), width=Inches(6))
                doc.add_paragraph()


        # Add Idle Event Durations chart
        doc.add_paragraph("Idle Event Durations:")
        fig_duration = go.Figure()
        fig_duration.add_trace(go.Bar(
            x=[f"Group {idx}" for idx in report_filtered_long_idles.index],
            y=report_filtered_long_idles['duration'],
            name='Duration (minutes)',
            marker_color='#ff6b6b'
        ))
        fig_duration.update_layout(title="Idle Event Durations by Group", xaxis_title="Idle Group", yaxis_title="Duration (minutes)")
        img_bytes_duration = pio.to_image(fig_duration, format="png")
        doc.add_picture(io.BytesIO(img_bytes_duration), width=Inches(6))
        doc.add_paragraph()

    else:
        doc.add_paragraph("No idle events found in the selected time range for this report.")


    # Add Hourly Performance Analysis table and chart
    doc.add_heading('Hourly Performance Analysis', level=1)
    doc.add_paragraph("Hourly Performance Statistics (based on current filtered data):")
    # Use filtered_df for hourly stats in the report for consistency with graphs
    hourly_stats_for_report_df = filtered_df.groupby('hour').agg({ 
        'speed': ['mean', 'max', 'min'],
        'is_running': 'mean',
        'quantity': lambda x: x.iloc[-1] - x.iloc[0] if len(x) > 1 else 0
    }).round(2)
    hourly_stats_for_report_df.columns = ['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Ratio', 'Production']
    hourly_stats_for_report_df['Uptime_Percent'] = (hourly_stats_for_report_df['Uptime_Ratio'] * 100).round(1)
    doc.add_paragraph(hourly_stats_for_report_df[['Avg_Speed', 'Max_Speed', 'Min_Speed', 'Uptime_Percent', 'Production']].to_string())

    doc.add_paragraph("Hourly Performance Chart (Average Speed and Uptime):")
    hourly_chart_data_for_report = hourly_stats_for_report_df[['Avg_Speed', 'Uptime_Percent']].copy()
    fig_hourly = create_multi_bar_chart(hourly_chart_data_for_report, ['Avg_Speed', 'Uptime_Percent'],
                                         "Hourly Performance - Average Speed and Uptime")
    img_bytes_hourly = pio.to_image(fig_hourly, format="png")
    doc.add_picture(io.BytesIO(img_bytes_hourly), width=Inches(6))
    doc.add_paragraph()

    # DeepSeek AI Analysis Results (Placeholder for actual content)
    doc.add_heading('DeepSeek AI Analysis Results', level=1)
    doc.add_paragraph("This section contains the detailed analysis generated by DeepSeek AI. Please run the AI analysis in the Streamlit app and manually add relevant sections here from the AI output displayed on the dashboard for a complete report.")
    doc.add_paragraph("Example: [Insert DeepSeek AI 'Full Dataset Analysis' here]")
    doc.add_paragraph("Example: [Insert DeepSeek AI 'Maintenance Recommendations' here]")
    doc.add_paragraph("Example: [Insert DeepSeek AI 'Production Optimization' here]")
    
    return doc

# Download Report Button
st.subheader("Generate Comprehensive Report")
if st.button("Download Report (DOCX)"):
    if 'df' in locals() and 'filtered_df' in locals():
        kpi_data_for_report = {
            'uptime': uptime_percent,
            'avg_speed': avg_speed,
            'total_production': total_production,
            'idle_events': len(long_idles) # This uses the overall long_idles
        }
        
        doc = create_comprehensive_report_with_graphs(df, filtered_df, min_speed, idle_time, kpi_data_for_report, selected_file)
        
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="Click to Download Report",
            data=bio.getvalue(),
            file_name=f"Machine_Analytics_Report_{selected_file.replace('.csv', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Please upload and analyze data first to generate a report.")